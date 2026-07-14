"""Render nstc-meeting-program-planning.md into an academically-formatted .docx.

A purpose-built Markdown renderer for this one meeting document. It is intentionally
small and specific — it understands only the subset of Markdown the source uses:
ATX headings (#/##/###), pipe tables, images, bold (**), inline code (`), a trailing
italic note, and single-line paragraphs.

Typography (Taiwan academic convention):
  - Title / headings : Microsoft JhengHei (微軟正黑體), bold
  - Body / tables    : PMingLiU (新細明體) for CJK, Times New Roman for Latin
  - A4, 2.5cm margins, 1.5 line spacing, justified body with a 2-char first-line indent
  - Bordered header cells with a light-blue fill; centred figures with 圖 captions;
    centred page numbers in the footer.

Usage (from anywhere):
    python docs/assessments/build_meeting_docx.py

Re-run any time the Markdown changes; the .docx is regenerated in place.
"""

from __future__ import annotations

import re
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = Path(__file__).resolve().parent
SRC_MD = BASE / "nstc-meeting-program-planning.md"
OUT_DOCX = BASE / "nstc-meeting-program-planning.docx"

# --- Typography constants -------------------------------------------------------
CJK_BODY = "PMingLiU"          # 新細明體
CJK_HEAD = "Microsoft JhengHei"  # 微軟正黑體
LATIN_BODY = "Times New Roman"
MONO = "Consolas"

INK = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0x1F, 0x3B, 0x63)   # deep blue for chapter headings
SUBTLE = RGBColor(0x59, 0x59, 0x59)   # grey for subtitle/caption
HEADER_FILL = "D9E2F3"                 # light blue table header

BODY_PT = 12
TABLE_PT = 10.5

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


# --- Low-level helpers ----------------------------------------------------------
def _set_fonts(run, *, latin, cjk, size=None, bold=None, italic=None, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_inline(par, text, *, size, cjk=CJK_BODY, latin=LATIN_BODY, color=INK, base_bold=False):
    """Add runs to a paragraph, honouring **bold** and `code` spans."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2])
            _set_fonts(r, latin=latin, cjk=cjk, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            _set_fonts(r, latin=MONO, cjk=cjk, size=(size - 0.5 if size else None), color=color)
        else:
            r = par.add_run(part)
            _set_fonts(r, latin=latin, cjk=cjk, size=size, bold=base_bold, color=color)


def _bottom_border(par, color="9DB3D0", size="6"):
    ppr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def _shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def _page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for tag, attrs, txt in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, "PAGE"),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        run = p.add_run()
        _set_fonts(run, latin=LATIN_BODY, cjk=CJK_BODY, size=10, color=SUBTLE)
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        if txt:
            el.text = txt
        run._r.append(el)


# --- Block renderers ------------------------------------------------------------
def _new_para(doc, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0.0, after=4.0,
              line=1.5, indent_chars=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if indent_chars:
        pf.first_line_indent = Pt(BODY_PT * indent_chars)
    return p


def render_title(doc, text):
    p = _new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6, line=1.3)
    _add_inline(p, text, size=20, cjk=CJK_HEAD, latin=CJK_HEAD, color=ACCENT)
    for r in p.runs:
        r.font.bold = True


def render_subtitle(doc, text):
    p = _new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10, line=1.2)
    _add_inline(p, text, size=11, cjk=CJK_HEAD, latin=CJK_HEAD, color=SUBTLE)


def render_h1(doc, text):
    p = _new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, before=16, after=8, line=1.2)
    p.paragraph_format.keep_with_next = True
    _add_inline(p, text, size=16, cjk=CJK_HEAD, latin=CJK_HEAD, color=ACCENT)
    for r in p.runs:
        r.font.bold = True
    _bottom_border(p)


def render_h2(doc, text):
    p = _new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=5, line=1.2)
    p.paragraph_format.keep_with_next = True
    _add_inline(p, text, size=13.5, cjk=CJK_HEAD, latin=CJK_HEAD, color=INK)
    for r in p.runs:
        r.font.bold = True


def render_body(doc, text):
    p = _new_para(doc, indent_chars=2)
    _add_inline(p, text, size=BODY_PT)


def render_note(doc, text):
    p = _new_para(doc, before=8, after=2, line=1.3)
    _add_inline(p, text, size=10.5, color=SUBTLE)
    for r in p.runs:
        r.font.italic = True


def render_image(doc, alt, path, fig_no):
    real = BASE / unquote(path)
    p = _new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2, line=1.0)
    run = p.add_run()
    run.add_picture(str(real), width=Cm(15))
    cap = _new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10, line=1.1)
    _add_inline(cap, f"圖 {fig_no}　{alt}", size=10, cjk=CJK_HEAD, latin=CJK_HEAD, color=SUBTLE)


def render_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=len(rows), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for c, cell_text in enumerate(header):
        cell = table.cell(0, c)
        _shade(cell, HEADER_FILL)
        par = cell.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_after = Pt(1)
        par.paragraph_format.line_spacing = 1.0
        _add_inline(par, cell_text, size=TABLE_PT, cjk=CJK_HEAD, latin=CJK_HEAD, base_bold=True)
    for r, row in enumerate(body, start=1):
        for c, cell_text in enumerate(row):
            cell = table.cell(r, c)
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(1)
            par.paragraph_format.line_spacing = 1.05
            _add_inline(par, cell_text, size=TABLE_PT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# --- Markdown parsing -----------------------------------------------------------
def split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def build():
    lines = SRC_MD.read_text(encoding="utf-8").splitlines()

    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)  # A4
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Cm(2.5))
    _page_number_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_BODY
    normal.font.size = Pt(BODY_PT)
    n_rpr = normal.element.get_or_add_rPr()
    n_fonts = n_rpr.find(qn("w:rFonts"))
    if n_fonts is None:
        n_fonts = OxmlElement("w:rFonts")
        n_rpr.append(n_fonts)
    n_fonts.set(qn("w:eastAsia"), CJK_BODY)

    fig_no = 0
    seen_title = False
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            render_title(doc, stripped[2:].strip())
            seen_title = True
            i += 1
            continue
        if stripped.startswith("### "):
            render_h2(doc, stripped[4:].strip())
            i += 1
            continue
        if stripped.startswith("## "):
            render_h1(doc, stripped[3:].strip())
            i += 1
            continue

        if stripped.startswith("!["):
            m = re.match(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)", stripped)
            if m:
                fig_no += 1
                render_image(doc, m.group("alt"), m.group("path"), fig_no)
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = [split_row(r) for r in block]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]  # drop separator row
            if rows:
                render_table(doc, rows)
            continue

        # subtitle line right after the title (the "会议…｜日期" meta line)
        if seen_title and "｜" in stripped and "日期" in stripped:
            render_subtitle(doc, stripped)
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            render_note(doc, stripped.strip("*").strip())
            i += 1
            continue

        render_body(doc, stripped)
        i += 1

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}  ({OUT_DOCX.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build()
