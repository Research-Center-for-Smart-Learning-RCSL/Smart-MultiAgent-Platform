"""Document parsers (R10.01 / R10.03).

Supports: pdf, docx, md, txt. Explicitly **no HTML**.

Path parsers enforce extraction budgets and may be run through
``parse_path_isolated`` so hostile PDF/DOCX expansion is confined to a
resource-limited child process. Legacy byte parsers remain for bounded callers.

SoC: parser functions are pure byte-in / str-out; they do not touch the
DB, MinIO, or Qdrant. Callers (RAG ingest, chat attachment extraction) wire
them together. Lives in shared_kernel because both the ``knowledge`` and
``conversation`` bounded contexts need it and contexts must not import each
other's infrastructure.
"""

from __future__ import annotations

import contextlib
import io
import mimetypes
import multiprocessing
import os
import shutil
import signal
import subprocess  # — invocation is fully controlled below
import tempfile
import zipfile
from collections.abc import Callable
from dataclasses import dataclass
from multiprocessing.connection import Connection
from multiprocessing.process import BaseProcess
from pathlib import Path

__all__ = [
    "MIME_TO_PARSER",
    "SUPPORTED_MIMES",
    "ExtractionLimits",
    "ParserError",
    "ResourceBudgetError",
    "normalise_mime",
    "parse_path",
    "parse_path_isolated",
    "parse_docx",
    "parse_markdown",
    "parse_pdf",
    "parse_plaintext",
]


def normalise_mime(raw: str, filename: str) -> str:
    """Prefer the client-supplied MIME; fall back to filename sniff.

    Shared by RAG and Knowledge Map ingest (both used to carry their own
    identical copy — collapsed here in code review, 2026-07-10).
    """
    # Strip MIME parameters (e.g. "; charset=utf-8") so downstream lookups
    # match the bare media type (M6).
    raw = raw.split(";")[0].strip()
    if raw and raw not in {"application/octet-stream", ""}:
        return raw
    guessed, _ = mimetypes.guess_type(filename)
    return guessed or "application/octet-stream"


class ParserError(RuntimeError):
    """Generic parser failure; mapped to `knowledge/ingest-failed` at the edge."""


class ResourceBudgetError(ParserError):
    """Extraction stopped at a deterministic resource boundary."""


@dataclass(frozen=True, slots=True)
class ExtractionLimits:
    extracted_utf8_bytes: int = 64 * 1024 * 1024
    estimated_tokens: int = 10_000_000
    pdf_pages: int = 5_000
    ocr_pages: int = 100
    docx_expanded_bytes: int = 256 * 1024 * 1024
    docx_compression_ratio: int = 100
    docx_entries: int = 10_000
    docx_single_entry_bytes: int = 64 * 1024 * 1024


DEFAULT_EXTRACTION_LIMITS = ExtractionLimits()
_PARSER_TIMEOUT_SECONDS = 300
_PARSER_MEMORY_BYTES = 1024 * 1024 * 1024


class _BoundedText:
    def __init__(self, limits: ExtractionLimits) -> None:
        self._limits = limits
        self._parts: list[str] = []
        self._utf8_bytes = 0
        self._characters = 0

    def append(self, value: str) -> None:
        encoded_size = len(value.encode("utf-8"))
        if self._utf8_bytes + encoded_size > self._limits.extracted_utf8_bytes:
            raise ResourceBudgetError("extracted_utf8_bytes")
        if (self._characters + len(value) + 3) // 4 > self._limits.estimated_tokens:
            raise ResourceBudgetError("estimated_tokens")
        self._parts.append(value)
        self._utf8_bytes += encoded_size
        self._characters += len(value)

    def append_section(self, value: str) -> None:
        self.append(("\n\n" if self._parts else "") + value)

    def finish(self) -> str:
        return "".join(self._parts)


def parse_path(
    path: Path,
    mime: str,
    *,
    limits: ExtractionLimits = DEFAULT_EXTRACTION_LIMITS,
) -> str:
    """Extract a spooled source while enforcing deterministic hard budgets."""
    if mime in {"text/plain", "text/markdown"}:
        return _parse_text_path(path, limits)
    if mime == "application/pdf":
        return _parse_pdf_path(path, limits)
    if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx_path(path, limits)
    raise ParserError(f"unsupported mime {mime!r}")


def _isolated_parse_worker(
    path: str,
    mime: str,
    limits: ExtractionLimits,
    connection: Connection,
) -> None:
    try:
        if os.name == "posix":
            import resource

            os.setsid()  # type: ignore[attr-defined]
            resource.setrlimit(  # type: ignore[attr-defined]
                resource.RLIMIT_AS,  # type: ignore[attr-defined]
                (_PARSER_MEMORY_BYTES, _PARSER_MEMORY_BYTES),
            )
            resource.setrlimit(  # type: ignore[attr-defined]
                resource.RLIMIT_CPU,  # type: ignore[attr-defined]
                (_PARSER_TIMEOUT_SECONDS, _PARSER_TIMEOUT_SECONDS),
            )
        connection.send(("ok", parse_path(Path(path), mime, limits=limits)))
    except ResourceBudgetError as exc:
        connection.send(("budget", str(exc)))
    except ParserError as exc:
        connection.send(("parser", str(exc)))
    except BaseException:
        connection.send(("parser", "document parser failed"))
    finally:
        connection.close()


def _terminate_process_tree(process: BaseProcess) -> None:
    if process.pid is None or not process.is_alive():
        return
    if os.name == "nt":
        subprocess.run(  # noqa: S603
            ["taskkill", "/PID", str(process.pid), "/T", "/F"],  # noqa: S607
            check=False,
            capture_output=True,
            timeout=10,
        )
    else:
        with contextlib.suppress(ProcessLookupError):
            os.killpg(process.pid, signal.SIGKILL)  # type: ignore[attr-defined]
    process.join(timeout=5)
    if process.is_alive():
        process.kill()
        process.join(timeout=5)


def parse_path_isolated(
    path: Path,
    mime: str,
    *,
    limits: ExtractionLimits = DEFAULT_EXTRACTION_LIMITS,
    timeout_seconds: int = _PARSER_TIMEOUT_SECONDS,
) -> str:
    """Parse in a bounded child process and return only validated text."""
    context = multiprocessing.get_context("spawn")
    parent, child = context.Pipe(duplex=False)
    process = context.Process(
        target=_isolated_parse_worker,
        args=(str(path), mime, limits, child),
        name="knowledge-document-parser",
    )
    process.start()
    child.close()
    try:
        if not parent.poll(timeout_seconds):
            _terminate_process_tree(process)
            raise ResourceBudgetError("parser_timeout")
        kind, payload = parent.recv()
    except EOFError as exc:
        raise ResourceBudgetError("parser_process_memory") from exc
    finally:
        parent.close()
        if process.is_alive():
            _terminate_process_tree(process)
        else:
            process.join(timeout=5)
    if kind == "ok":
        return str(payload)
    if kind == "budget":
        raise ResourceBudgetError(str(payload))
    raise ParserError(str(payload))


def _parse_text_path(path: Path, limits: ExtractionLimits) -> str:
    import codecs

    decoder = codecs.getincrementaldecoder("utf-8")(errors="replace")
    output = _BoundedText(limits)
    with path.open("rb") as source:
        while chunk := source.read(1024 * 1024):
            output.append(decoder.decode(chunk))
    output.append(decoder.decode(b"", final=True))
    return output.finish()


def _parse_pdf_path(path: Path, limits: ExtractionLimits) -> str:
    try:
        import pypdf
    except ImportError as exc:  # pragma: no cover
        raise ParserError("pypdf not installed") from exc
    try:
        reader = pypdf.PdfReader(path)
        if len(reader.pages) > limits.pdf_pages:
            raise ResourceBudgetError("pdf_pages")
        output = _BoundedText(limits)
        for page in reader.pages:
            try:
                text = (page.extract_text() or "").strip()
            except Exception:
                text = ""
            if text:
                output.append_section(text)
    except ResourceBudgetError:
        raise
    except Exception as exc:
        raise ParserError(f"pdf parse failed: {exc}") from exc
    extracted = output.finish()
    if extracted:
        return extracted
    if shutil.which("tesseract") is None:
        return ""
    if len(reader.pages) > limits.ocr_pages:
        raise ResourceBudgetError("ocr_pages")
    return _tesseract_ocr_path(path, limits, page_count=len(reader.pages))


def _tesseract_ocr_path(
    path: Path,
    limits: ExtractionLimits,
    *,
    page_count: int,
) -> str:
    if shutil.which("pdftoppm") is None:
        return ""
    with tempfile.TemporaryDirectory() as tmpdir:
        base = Path(tmpdir)
        output = _BoundedText(limits)
        for page_number in range(1, page_count + 1):
            image = base / "page.png"
            try:
                subprocess.run(  # noqa: S603
                    [  # noqa: S607
                        "pdftoppm",
                        "-f",
                        str(page_number),
                        "-l",
                        str(page_number),
                        "-singlefile",
                        "-r",
                        "200",
                        "-png",
                        str(path),
                        str(base / "page"),
                    ],
                    check=True,
                    capture_output=True,
                    timeout=120,
                )
                result = subprocess.run(  # noqa: S603
                    ["tesseract", str(image), "-"],  # noqa: S607
                    check=True,
                    capture_output=True,
                    timeout=120,
                )
            except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
                continue
            text = result.stdout.decode("utf-8", errors="replace").strip()
            if text:
                output.append_section(text)
            image.unlink(missing_ok=True)
        return output.finish()


def _parse_docx_path(path: Path, limits: ExtractionLimits) -> str:
    try:
        with zipfile.ZipFile(path) as archive:
            entries = archive.infolist()
            if len(entries) > limits.docx_entries:
                raise ResourceBudgetError("docx_entries")
            expanded = 0
            compressed = 0
            for entry in entries:
                if entry.flag_bits & 0x1:
                    raise ParserError("encrypted docx is not supported")
                if entry.file_size > limits.docx_single_entry_bytes:
                    raise ResourceBudgetError("docx_single_entry_bytes")
                compressed += entry.compress_size
                entry_expanded = 0
                with archive.open(entry) as source:
                    while chunk := source.read(1024 * 1024):
                        entry_expanded += len(chunk)
                        expanded += len(chunk)
                        if entry_expanded > limits.docx_single_entry_bytes:
                            raise ResourceBudgetError("docx_single_entry_bytes")
                        if expanded > limits.docx_expanded_bytes:
                            raise ResourceBudgetError("docx_expanded_bytes")
            if expanded > limits.docx_compression_ratio * max(compressed, 1):
                raise ResourceBudgetError("docx_compression_ratio")
    except (ParserError, ResourceBudgetError):
        raise
    except (OSError, zipfile.BadZipFile) as exc:
        raise ParserError(f"docx parse failed: {exc}") from exc
    try:
        import docx

        # str(), not the Path: python-docx opens either, but its stubs accept
        # only `str | IO[bytes] | None`.
        document = docx.Document(str(path))
    except Exception as exc:
        raise ParserError(f"docx parse failed: {exc}") from exc
    output = _BoundedText(limits)
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            output.append_section(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    output.append_section(text)
    return output.finish()


def parse_plaintext(data: bytes) -> str:
    """UTF-8 decode with `errors="replace"` so mis-labelled files never crash
    the ingest path. Invalid sequences become U+FFFD; the chunker will still
    produce usable chunks for the downstream embedder."""
    return data.decode("utf-8", errors="replace")


def parse_markdown(data: bytes) -> str:
    """Markdown is forwarded verbatim — we do NOT render to HTML because
    downstream chunkers work better on source markdown than on rendered
    text (headings stay intact, lists stay list-like)."""
    return parse_plaintext(data)


def parse_pdf(data: bytes) -> str:
    """Extract text via pypdf; fall back to tesseract if the text layer is
    empty AND tesseract is installed.

    We import pypdf lazily so a missing install does not crash module load.
    """
    try:
        import pypdf
    except ImportError as exc:  # pragma: no cover — pypdf is in the runtime deps
        raise ParserError("pypdf not installed") from exc

    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
    except Exception as exc:  # — pypdf raises many error types
        raise ParserError(f"pdf parse failed: {exc}") from exc

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    text = "\n\n".join(p.strip() for p in pages if p.strip())
    if text:
        return text

    # Text layer empty — try OCR if tesseract is on PATH.
    if shutil.which("tesseract") is None:
        return ""
    return _tesseract_ocr(data)


def _tesseract_ocr(pdf_bytes: bytes) -> str:
    """Shell out to `tesseract` page-by-page. Returns "" on any failure.

    We keep this entirely subprocess-based (no pytesseract dependency) so
    the OCR path does not add a wheel to the runtime environment. Each
    page is rasterised via `pdftoppm` which is part of poppler-utils on
    the same container tesseract comes from.
    """
    if shutil.which("pdftoppm") is None:
        return ""
    with tempfile.TemporaryDirectory() as tmpdir:
        base = Path(tmpdir)
        pdf_path = base / "src.pdf"
        pdf_path.write_bytes(pdf_bytes)
        try:
            subprocess.run(  # noqa: S603
                ["pdftoppm", "-r", "200", "-png", str(pdf_path), str(base / "page")],  # noqa: S607
                check=True,
                capture_output=True,
                timeout=120,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
            return ""
        pages_text: list[str] = []
        for img in sorted(base.glob("page-*.png")):
            try:
                r = subprocess.run(  # noqa: S603
                    ["tesseract", str(img), "-"],  # noqa: S607
                    check=True,
                    capture_output=True,
                    timeout=120,
                )
                pages_text.append(r.stdout.decode("utf-8", errors="replace"))
            except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
                continue
        return "\n\n".join(t.strip() for t in pages_text if t.strip())


def parse_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ParserError("python-docx not installed") from exc

    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ParserError(f"docx parse failed: {exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Include table cells — agents use docx for spec tables constantly.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    return "\n\n".join(parts)


# MIME dispatch table (R10.03 — no HTML).
MIME_TO_PARSER: dict[str, Callable[[bytes], str]] = {
    "application/pdf": parse_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": parse_docx,
    "text/markdown": parse_markdown,
    "text/plain": parse_plaintext,
}

SUPPORTED_MIMES: frozenset[str] = frozenset(MIME_TO_PARSER.keys())
