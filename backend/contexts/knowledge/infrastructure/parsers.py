"""Document parsers (R10.01 / R10.03).

Supports: pdf, docx, md, txt. Explicitly **no HTML**.

Each parser accepts the raw file bytes and returns a flat UTF-8 string
(whitespace-normalised, preserving paragraph breaks). Optional OCR on PDFs
runs only if ``tesseract`` is present on PATH and the extracted text is
empty — we never pay the OCR cost on text-bearing PDFs.

SoC: parser functions are pure byte-in / str-out; they do not touch the
DB, MinIO, or Qdrant. The ingest service wires them together.
"""

from __future__ import annotations

import io
import shutil
import subprocess  # — invocation is fully controlled below
import tempfile
from collections.abc import Callable
from pathlib import Path

__all__ = [
    "MIME_TO_PARSER",
    "SUPPORTED_MIMES",
    "ParserError",
    "parse_docx",
    "parse_markdown",
    "parse_pdf",
    "parse_plaintext",
]


class ParserError(RuntimeError):
    """Generic parser failure; mapped to `knowledge/ingest-failed` at the edge."""


def parse_plaintext(data: bytes) -> str:
    """UTF-8 decode with `errors="replace"` so mis-labelled files never crash
    the ingest path. Invalid sequences become U+FFFD; the chunker will still
    produce usable chunks for the downstream embedder."""
    return data.decode("utf-8", errors="replace")


def parse_markdown(data: bytes) -> str:
    """Markdown is forwarded verbatim — we do NOT render to HTML because
    downstream chunkers work better on source markdown than on rendered
    text (headings stay intact, lists stay list-like)."""
    return parse_plaintext(data)


def parse_pdf(data: bytes) -> str:
    """Extract text via pypdf; fall back to tesseract if the text layer is
    empty AND tesseract is installed.

    We import pypdf lazily so a missing install does not crash module load.
    """
    try:
        import pypdf
    except ImportError as exc:  # pragma: no cover — pypdf is in the runtime deps
        raise ParserError("pypdf not installed") from exc

    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
    except Exception as exc:  # — pypdf raises many error types
        raise ParserError(f"pdf parse failed: {exc}") from exc

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    text = "\n\n".join(p.strip() for p in pages if p.strip())
    if text:
        return text

    # Text layer empty — try OCR if tesseract is on PATH.
    if shutil.which("tesseract") is None:
        return ""
    return _tesseract_ocr(data)


def _tesseract_ocr(pdf_bytes: bytes) -> str:
    """Shell out to `tesseract` page-by-page. Returns "" on any failure.

    We keep this entirely subprocess-based (no pytesseract dependency) so
    the OCR path does not add a wheel to the runtime environment. Each
    page is rasterised via `pdftoppm` which is part of poppler-utils on
    the same container tesseract comes from.
    """
    if shutil.which("pdftoppm") is None:
        return ""
    with tempfile.TemporaryDirectory() as tmpdir:
        base = Path(tmpdir)
        pdf_path = base / "src.pdf"
        pdf_path.write_bytes(pdf_bytes)
        try:
            subprocess.run(  # noqa: S603
                ["pdftoppm", "-r", "200", "-png", str(pdf_path), str(base / "page")],  # noqa: S607
                check=True,
                capture_output=True,
                timeout=120,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
            return ""
        pages_text: list[str] = []
        for img in sorted(base.glob("page-*.png")):
            try:
                r = subprocess.run(  # noqa: S603
                    ["tesseract", str(img), "-"],  # noqa: S607
                    check=True,
                    capture_output=True,
                    timeout=120,
                )
                pages_text.append(r.stdout.decode("utf-8", errors="replace"))
            except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
                continue
        return "\n\n".join(t.strip() for t in pages_text if t.strip())


def parse_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ParserError("python-docx not installed") from exc

    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ParserError(f"docx parse failed: {exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Include table cells — agents use docx for spec tables constantly.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    return "\n\n".join(parts)


# MIME dispatch table (R10.03 — no HTML).
MIME_TO_PARSER: dict[str, Callable[[bytes], str]] = {
    "application/pdf": parse_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": parse_docx,
    "text/markdown": parse_markdown,
    "text/plain": parse_plaintext,
}

SUPPORTED_MIMES: frozenset[str] = frozenset(MIME_TO_PARSER.keys())
